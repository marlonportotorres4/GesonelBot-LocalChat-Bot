"""
Processamento de documentos do GesonelBot

Este módulo é responsável por processar os documentos carregados pelos usuários,
extrair seu conteúdo e preparar os dados para indexação.
"""
import os
import mimetypes
import hashlib
import logging
from typing import List, Dict, Optional, Tuple, Any
from datetime import datetime

# Importar componentes do LangChain para processamento de texto
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain.docstore.document import Document

# Importar configurações e componentes do GesonelBot
from gesonelbot.config.settings import (
    DOCS_DIR, 
    VECTORSTORE_DIR, 
    CHUNK_SIZE, 
    CHUNK_OVERLAP
)
from gesonelbot.core.embeddings_manager import embeddings_manager

# Configurar logging
logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(name)s - %(levelname)s - %(message)s')
logger = logging.getLogger(__name__)

# Configurar tipos MIME suportados
SUPPORTED_MIME_TYPES = {
    'text/plain': '.txt',
    'application/pdf': '.pdf',
    'application/vnd.openxmlformats-officedocument.wordprocessingml.document': '.docx'
}

# Importações para processamento de documentos
try:
    import docx
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    logger.warning("Biblioteca python-docx não encontrada. O processamento de arquivos DOCX não estará disponível.")

try:
    import pypdf
    PDF_AVAILABLE = True
except ImportError:
    PDF_AVAILABLE = False
    logger.warning("Biblioteca pypdf não encontrada. O processamento de arquivos PDF não estará disponível.")

def extract_text_from_txt(file_path: str) -> str:
    """
    Extrai texto de um arquivo TXT.
    
    Args:
        file_path: Caminho para o arquivo TXT
        
    Returns:
        Texto extraído do arquivo
    """
    try:
        # Primeiro tenta com codificação UTF-8
        with open(file_path, 'r', encoding='utf-8') as file:
            return file.read()
    except UnicodeDecodeError:
        # Se falhar, tenta com codificação alternativa
        logger.info(f"Falha ao decodificar {file_path} com UTF-8, tentando latin-1")
        with open(file_path, 'r', encoding='latin-1') as file:
            return file.read()

def extract_text_from_docx(file_path: str) -> str:
    """
    Extrai texto de um arquivo DOCX.
    
    Args:
        file_path: Caminho para o arquivo DOCX
        
    Returns:
        Texto extraído do arquivo ou mensagem de erro
    """
    if not DOCX_AVAILABLE:
        return f"[Biblioteca python-docx não instalada] Não foi possível extrair o texto de {os.path.basename(file_path)}"
    
    try:
        # Carregar o documento
        doc = docx.Document(file_path)
        
        # Extrair texto de todos os parágrafos
        full_text = []
        for para in doc.paragraphs:
            full_text.append(para.text)
            
        # Extrair texto de tabelas
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    full_text.append(cell.text)
        
        # Juntar todo o texto com quebras de linha
        return '\n'.join(full_text)
    except Exception as e:
        logger.error(f"Erro ao processar DOCX {file_path}: {str(e)}")
        return f"[Erro ao processar DOCX: {str(e)}] Não foi possível extrair o texto de {os.path.basename(file_path)}"

def extract_text_from_pdf(file_path: str) -> str:
    """
    Extrai texto de um arquivo PDF.
    
    Args:
        file_path: Caminho para o arquivo PDF
        
    Returns:
        Texto extraído do arquivo ou mensagem de erro
    """
    if not PDF_AVAILABLE:
        return f"[Biblioteca pypdf não instalada] Não foi possível extrair o texto de {os.path.basename(file_path)}"
    
    try:
        text = []
        # Abrir o PDF
        with open(file_path, 'rb') as file:
            reader = pypdf.PdfReader(file)
            
            # Extrair texto de cada página
            for page_num in range(len(reader.pages)):
                page = reader.pages[page_num]
                text.append(page.extract_text())
        
        # Juntar todo o texto com quebras de linha
        return '\n'.join(text)
    except Exception as e:
        logger.error(f"Erro ao processar PDF {file_path}: {str(e)}")
        return f"[Erro ao processar PDF: {str(e)}] Não foi possível extrair o texto de {os.path.basename(file_path)}"

def validate_file(file_path: str) -> Tuple[bool, str]:
    """
    Valida se o arquivo é suportado e está em bom estado.
    
    Args:
        file_path: Caminho para o arquivo
        
    Returns:
        Tupla (é_válido, mensagem)
    """
    # Verificar se o arquivo existe
    if not os.path.exists(file_path):
        return False, "Arquivo não encontrado"
    
    # Verificar se é um arquivo
    if not os.path.isfile(file_path):
        return False, "O caminho especificado não é um arquivo"
    
    # Verificar tamanho do arquivo (máximo 20MB)
    if os.path.getsize(file_path) > 20 * 1024 * 1024:
        return False, "Arquivo muito grande (máximo 20MB)"
    
    # Verificar tipo MIME
    mime_type, _ = mimetypes.guess_type(file_path)
    if mime_type not in SUPPORTED_MIME_TYPES:
        return False, f"Tipo de arquivo não suportado: {mime_type}"
    
    return True, "Arquivo válido"

def get_file_metadata(file_path: str) ->  Dict[str, str]:
    """
    Obtém metadados do arquivo.
    
    Args:
        file_path: Caminho para o arquivo
        
    Returns:
        Dicionário com metadados do arquivo
    """
    file_stat = os.stat(file_path)
    
    # Calcular hash do arquivo para identificação única
    file_hash = hashlib.md5()
    with open(file_path, 'rb') as f:
        for chunk in iter(lambda: f.read(4096), b''):
            file_hash.update(chunk)
    
    return {
        "file_name": os.path.basename(file_path),
        "file_path": file_path,  # Caminho completo para referência
        "file_size": str(file_stat.st_size),
        "created_at": datetime.fromtimestamp(file_stat.st_ctime).isoformat(),
        "modified_at": datetime.fromtimestamp(file_stat.st_mtime).isoformat(),
        "file_hash": file_hash.hexdigest(),
        "mime_type": mimetypes.guess_type(file_path)[0],
        "source": f"uploaded:{os.path.basename(file_path)}"  # Identificador de fonte para citações
    }

def process_document(file_path: str) -> Dict[str, Any]:
    """
    Processa um único documento e extrai seu texto.
    
    Args:
        file_path: Caminho para o arquivo a ser processado
        
    Returns:
        Dicionário com informações sobre o documento processado
    """
    # Validar arquivo
    is_valid, message = validate_file(file_path)
    if not is_valid:
        logger.warning(f"Arquivo inválido: {file_path} - {message}")
        return {
            "status": "error",
            "file_name": os.path.basename(file_path),
            "message": message,
            "text": ""
        }
    
    # Obter metadados
    metadata = get_file_metadata(file_path)
    logger.info(f"Processando documento: {metadata['file_name']}")
    
    # Determinar o tipo de arquivo
    file_extension = os.path.splitext(file_path)[1].lower()
    
    # Extrair texto com base no tipo de arquivo
    try:
        if file_extension == '.txt':
            text = extract_text_from_txt(file_path)
        elif file_extension == '.docx':
            text = extract_text_from_docx(file_path)
        elif file_extension == '.pdf':
            text = extract_text_from_pdf(file_path)
        else:
            logger.warning(f"Formato não suportado: {file_extension}")
            return {
                "status": "error",
                "file_name": metadata["file_name"],
                "message": f"Formato não suportado: {file_extension}",
                "text": ""
            }
            
        # Verificar se o texto foi extraído com sucesso
        if not text or text.startswith("[Erro") or text.startswith("[Biblioteca"):
            logger.warning(f"Falha na extração de texto: {metadata['file_name']}")
            return {
                "status": "error",
                "file_name": metadata["file_name"],
                "message": f"Falha na extração de texto: {text if text else 'Texto vazio'}",
                "text": ""
            }
            
        # Retornar documento processado com sucesso
        logger.info(f"Documento processado com sucesso: {metadata['file_name']} ({len(text)} caracteres)")
        return {
            "status": "success",
            "file_name": metadata["file_name"],
            "message": "Documento processado com sucesso",
            "text": text,
            "metadata": metadata
        }
        
    except Exception as e:
        # Retornar erro em caso de falha
        logger.error(f"Erro ao processar documento {file_path}: {str(e)}")
        return {
            "status": "error",
            "file_name": metadata["file_name"],
            "message": f"Erro ao processar documento: {str(e)}",
            "text": "",
            "metadata": metadata
        }

def process_documents(file_paths: List[str]) -> Dict[str, object]:
    """
    Processa uma lista de documentos e prepara para indexação.
    
    Args:
        file_paths: Lista de caminhos para os arquivos a serem processados
        
    Returns:
        Dicionário com informações sobre o processamento
    """
    results = {
        "success_count": 0,
        "error_count": 0,
        "processed_files": [],
        "errors": []
    }
    
    # Processar cada arquivo
    for file_path in file_paths:
        result = process_document(file_path)
        
        # Registrar resultado
        if result["status"] == "success":
            results["success_count"] += 1
            results["processed_files"].append(result)  # Manter o resultado completo
        else:
            results["error_count"] += 1
            results["errors"].append(result)  # Manter o resultado completo para debug
    
    return results

def split_text_into_chunks(text: str, metadata: Dict[str, str]) -> List[Document]:
    """
    Divide o texto em chunks menores para processamento e indexação.
    
    Args:
        text: Texto a ser dividido
        metadata: Metadados do documento
        
    Returns:
        Lista de objetos Document do LangChain
    """
    if not text or len(text.strip()) == 0:
        logger.warning(f"Texto vazio para o documento {metadata.get('file_name', 'desconhecido')}")
        return []
    
    try:
        # Criar um divisor de texto que respeita parágrafos e frases
        # Usamos RecursiveCharacterTextSplitter para manter a estrutura semântica
        text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=CHUNK_SIZE,
            chunk_overlap=CHUNK_OVERLAP,
            length_function=len,
            separators=["\n\n", "\n", ". ", "! ", "? ", ";", ",", " ", ""]
        )
        
        # Dividir o texto em chunks
        chunks = text_splitter.create_documents(
            texts=[text],
            metadatas=[metadata]
        )
        
        logger.info(f"Documento '{metadata.get('file_name')}' dividido em {len(chunks)} chunks")
        
        return chunks
    except Exception as e:
        logger.error(f"Erro ao dividir texto em chunks: {str(e)}")
        # Fallback: criar um único documento se a divisão falhar
        return [Document(page_content=text, metadata=metadata)]

def get_processed_documents_info() -> List[Dict[str, str]]:
    """
    Retorna informações sobre todos os documentos processados.
    
    Returns:
        Lista de dicionários com informações sobre os documentos
    """
    if not os.path.exists(DOCS_DIR):
        logger.warning(f"Diretório de upload não existe: {DOCS_DIR}")
        return []
    
    documents_info = []
    
    try:
        for filename in os.listdir(DOCS_DIR):
            file_path = os.path.join(DOCS_DIR, filename)
            if os.path.isfile(file_path):
                # Verificar extensão
                _, ext = os.path.splitext(filename)
                if ext.lower() in ['.pdf', '.docx', '.txt']:
                    # Obter metadados básicos
                    stat = os.stat(file_path)
                    documents_info.append({
                        "file_name": filename,
                        "file_size": str(stat.st_size),
                        "created_at": datetime.fromtimestamp(stat.st_ctime).isoformat(),
                        "modified_at": datetime.fromtimestamp(stat.st_mtime).isoformat(),
                        "file_path": file_path,
                        "file_type": ext.lower()
                    })
        
        logger.info(f"Encontrados {len(documents_info)} documentos no diretório de upload")
    except Exception as e:
        logger.error(f"Erro ao obter informações dos documentos: {str(e)}")
    
    return documents_info

def get_total_upload_usage() -> int:
    """
    Calcula o uso total de armazenamento no diretório de upload.
    
    Returns:
        Tamanho total em bytes
    """
    if not os.path.exists(DOCS_DIR):
        return 0
    
    total_size = 0
    
    try:
        for filename in os.listdir(DOCS_DIR):
            file_path = os.path.join(DOCS_DIR, filename)
            if os.path.isfile(file_path):
                total_size += os.path.getsize(file_path)
        
        logger.info(f"Uso total do diretório de upload: {total_size/1024/1024:.2f}MB")
    except Exception as e:
        logger.error(f"Erro ao calcular uso total: {str(e)}")
    
    return total_size

def ingest_documents(file_paths: Optional[List[str]] = None) -> Dict[str, object]:
    """
    Processa documentos e os adiciona ao banco de dados vetorial.
    
    Args:
        file_paths: Lista opcional de caminhos para os arquivos a processar.
                    Se None, processa todos os arquivos no diretório de upload.
        
    Returns:
        Dicionário com informações sobre o processamento
    """
    # Se não foram especificados arquivos, usar todos do diretório de upload
    if file_paths is None:
        logger.info("Nenhum caminho específico fornecido, processando todos os documentos no diretório de upload")
        if not os.path.exists(DOCS_DIR):
            logger.warning(f"Diretório de upload não encontrado: {DOCS_DIR}")
            return {
                "success_count": 0,
                "error_count": 0,
                "message": f"Diretório de upload não encontrado: {DOCS_DIR}"
            }
        
        # Listar todos os arquivos com extensões suportadas
        file_paths = [
            os.path.join(DOCS_DIR, f) 
            for f in os.listdir(DOCS_DIR) 
            if os.path.isfile(os.path.join(DOCS_DIR, f)) and 
            os.path.splitext(f)[1].lower() in ['.pdf', '.docx', '.txt']
        ]
        
        logger.info(f"Encontrados {len(file_paths)} documentos para processamento")
    else:
        logger.info(f"Processando {len(file_paths)} documentos específicos")
    
    # Processar os documentos
    results = process_documents(file_paths)
    
    # Preparar documentos para indexação
    all_chunks = []
    
    # Processar cada documento em chunks para o banco de dados vetorial
    for doc_result in results["processed_files"]:
        try:
            # Dividir o texto em chunks
            chunks = split_text_into_chunks(
                doc_result["text"], 
                doc_result["metadata"]
            )
            
            if chunks:
                all_chunks.extend(chunks)
                logger.info(f"Adicionados {len(chunks)} chunks do documento '{doc_result['file_name']}'")
            else:
                logger.warning(f"Nenhum chunk gerado para o documento '{doc_result['file_name']}'")
        except Exception as e:
            logger.error(f"Erro ao processar chunks para '{doc_result['file_name']}': {str(e)}")
            results["error_count"] += 1
    
    # Adicionar os documentos ao banco de dados vetorial se houver chunks
    if all_chunks:
        try:
            # Usar o gerenciador de embeddings para adicionar documentos
            success = embeddings_manager.add_documents(all_chunks)
            if success:
                logger.info(f"Adicionados {len(all_chunks)} chunks ao banco de dados vetorial")
                results["vectorstore_chunks"] = len(all_chunks)
            else:
                logger.error("Falha ao adicionar documentos ao banco de dados vetorial")
                results["vectorstore_error"] = "Falha ao adicionar documentos ao banco de dados vetorial"
        except Exception as e:
            logger.error(f"Erro ao indexar documentos: {str(e)}")
            results["vectorstore_error"] = str(e)
    else:
        logger.warning("Nenhum chunk para adicionar ao banco de dados vetorial")
        results["vectorstore_chunks"] = 0
    
    # Adicionar informações resumidas
    results["summary"] = f"Processados {results['success_count']} documentos com sucesso, {results['error_count']} erros."
    if "vectorstore_chunks" in results:
        results["summary"] += f" {results['vectorstore_chunks']} chunks adicionados ao banco de dados vetorial."
    
    return results 